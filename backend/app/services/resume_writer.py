import re
import io
from typing import Optional
from docx import Document
from docx.shared import Pt
from app.services.style_extractor import ResumeStyle

SECTION_HEADERS = {"SUMMARY", "EXPERIENCE", "SKILLS", "EDUCATION"}


def _add_runs(paragraph, text: str, font_name: str, font_size_pt: int, default_bold: bool = False):
    """
    Write text into a paragraph, converting **bold** markers into real bold runs.
    Segments without markers inherit default_bold.
    """
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = default_bold
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)


class ResumeWriter:
    def create_docx(self, text: str, style: Optional[ResumeStyle] = None) -> bytes:
        """Build the .docx and return raw bytes — no local file written."""
        body_font   = style.body_font   if style else "Calibri"
        body_size   = style.body_size   if style else 11
        header_font = style.header_font if style else "Calibri"
        header_size = style.header_size if style else 14

        doc = Document()
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.strip('*').strip() in SECTION_HEADERS:
                clean = stripped.strip('*').strip()
                p = doc.add_paragraph()
                _add_runs(p, clean, header_font, header_size, default_bold=True)
            elif stripped.startswith("-"):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, stripped[1:].strip(), body_font, body_size)
            else:
                p = doc.add_paragraph()
                _add_runs(p, stripped, body_font, body_size)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
